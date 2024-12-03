import os
from docx import Document

# Specify the folder containing the .docx files
folder_path = "/Users/jbenno/temp/a"
output_file = "/Users/jbenno/temp/merged_document.docx"

# Create a new Word document for the merged content
merged_document = Document()

# Loop through all .docx files in the folder
for file_name in sorted(os.listdir(folder_path)):
    if file_name.endswith(".docx"):
        file_path = os.path.join(folder_path, file_name)
        
        # Open each Word document
        current_document = Document(file_path)
        
        # Append the content of each document to the merged document
        for paragraph in current_document.paragraphs:
            merged_document.add_paragraph(paragraph.text)
        
        # Add a page break after each document (optional)
        merged_document.add_page_break()

# Save the merged document
merged_document.save(output_file)

print(f"Merged document saved at: {output_file}")
